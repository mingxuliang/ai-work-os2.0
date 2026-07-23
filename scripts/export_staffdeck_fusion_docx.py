# -*- coding: utf-8 -*-
"""Export StaffDeck fusion development document to Word."""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]


def set_run_font(run, size=11, bold=False, color=None):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Microsoft YaHei"
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:eastAsia"), "微软雅黑")
    if color is not None:
        run.font.color.rgb = color


def h(doc, text, level=1):
    para = doc.add_heading(text, level=level)
    size = {1: 18, 2: 14, 3: 12}.get(level, 11)
    for run in para.runs:
        set_run_font(run, size=size, bold=True)


def p(doc, text, bold=False):
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_run_font(run, bold=bold)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.35


def bullets(doc, items):
    for item in items:
        para = doc.add_paragraph(style="List Bullet")
        run = para.add_run(item)
        set_run_font(run)
        para.paragraph_format.space_after = Pt(2)


def table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = header
        for para in cell.paragraphs:
            for run in para.runs:
                set_run_font(run, size=10, bold=True)
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = t.rows[r_idx + 1].cells[c_idx]
            cell.text = str(value)
            for para in cell.paragraphs:
                for run in para.runs:
                    set_run_font(run, size=10)
    doc.add_paragraph()


def main() -> None:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.4)
        section.right_margin = Cm(2.4)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("AIWork-OS × StaffDeck 能力融合详细开发文档")
    font(run, size=20, bold=True)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(
        f"日期：{date.today().isoformat()}  |  状态：可立项（按一期→二期→三期落地）"
    )
    font(run, size=10, color=RGBColor(0x66, 0x66, 0x66))

    h(doc, "一、融合结论", 1)
    p(
        doc,
        "可以融合。采用“能力吸收 + 自研落地”，不要整仓合并 StaffDeck 源码（AGPL-3.0）。",
        bold=True,
    )
    bullets(
        doc,
        [
            "一期：数字员工壳与市场（Agent 产品化包装）",
            "二期：状态机 SOP 引擎（流程编排、人工门、时间线回放）",
            "三期：文档结构检索 + 运行中接管与反馈闭环",
        ],
    )

    h(doc, "二、StaffDeck 核心功能", 1)
    table(
        doc,
        ["能力域", "功能要点"],
        [
            ["数字员工", "岗位、工号、能力画像、工作记录、权限隔离、发布复用"],
            ["SOP 状态机", "NL 生成 SOP、状态机执行、多流程切换、可视化编辑、版本分支"],
            ["结构感知知识库", "章/节/页索引、知识桶、引用、检索调试"],
            ["自主执行", "HTTP API、MCP、定时任务、长记忆、完整轨迹"],
            ["人机协同", "人工接管、用户反馈、反馈分析改进闭环"],
        ],
    )

    h(doc, "三、与 AIWork-OS 对照结论", 1)
    p(
        doc,
        "工具执行、定时、MCP、治理底座可直接复用；需要补齐的是数字员工产品层、SOP 状态机、结构检索与接管反馈。",
        bold=True,
    )

    h(doc, "四、一期详细设计（先做）", 1)
    p(doc, "一期只做“壳”，不改内核执行逻辑。", bold=True)
    p(
        doc,
        "数据表 digital_employees：employee_no、name、title、description、agent_id、"
        "owner_user_id、department_id、status、visibility、capability_json、tags_json、published_at。",
    )
    p(
        doc,
        "API 前缀 /api/digital-employees：列表、创建、详情、更新、发布、取消发布、克隆、统计、画廊。",
    )
    p(
        doc,
        "前端新增 DigitalEmployeeGallery、DigitalEmployeeProfile，并从 Workbench/导航进入；"
        "启用后跳转现有 Chat。",
    )
    p(doc, "验收：创建绑定、发布可见、权限隔离、启用进 Chat、API 单测通过。")

    h(doc, "五、二期详细设计：SOP 状态机", 1)
    p(
        doc,
        "SOP Runner 作为编排层，不改上游 Loop。节点类型：start / llm / tool / human_gate / "
        "branch / end。运行产生 SOPRun + SOPRunEvent，支持 resume/cancel 与时间线回放。",
    )
    p(
        doc,
        "MVP 范围：线性流程 + 单层分支 + 人工确认门；版本 draft/published；可绑定到数字员工。",
        bold=True,
    )

    h(doc, "六、三期详细设计", 1)
    p(
        doc,
        "在现有 RAG 上增加 outline_path 等结构字段；提供结构检索与调试 API；"
        "Chat 引用卡片。HITL：运行中接管、统一待办、反馈入库到员工档案。",
    )

    h(doc, "七、质量、合规与里程碑", 1)
    bullets(
        doc,
        [
            "开关：AIWORK_DIGITAL_EMPLOYEE / AIWORK_SOP_ENGINE / AIWORK_KNOWLEDGE_STRUCTURE",
            "测试：单测 + API 契约 + 前端 smoke + 回归 Chat/素材库/RAG",
            "合规：禁止 AGPL 源码并入；对外表述为自研能力吸收",
            "里程碑：M1 评审 → M2 一期上线 → M3 SOP MVP → M4 结构检索与接管",
            "粗估总工期约 9–14 人周",
        ],
    )

    h(doc, "九、立项后 3 天行动项", 1)
    bullets(
        doc,
        [
            "确认一期画廊字段与信息架构",
            "建立 digital_employees 表与 CRUD",
            "前端导航壳 + Gallery mock",
            "书面确认不引入 StaffDeck AGPL 源码",
            "SOP 画布技术选型评审",
        ],
    )

    out_md = ROOT / "docs" / "digital-employee" / "StaffDeck能力融合-详细开发文档.docx"
    desk = Path.home() / "Desktop" / "StaffDeck能力融合-详细开发文档.docx"
    doc.save(str(out1))
    doc.save(str(out2))
    print("Saved:")
    print(out1)
    print(out2)


if __name__ == "__main__":
    main()
